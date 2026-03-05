#!/usr/bin/env python3
"""
Medical Literature Monitor - Streamlit Demo App

Demo UI for the multi-agent literature monitoring system. Reuses all agent
logic from agents/ (SearchAgent, AnalysisAgent, ReportAgent). Keeps main.py
unchanged for CLI usage.

Usage:
    streamlit run app_demo.py
    python -m streamlit run app_demo.py   # if streamlit not on PATH

Environment:
    ANTHROPIC_API_KEY must be set (required for AnalysisAgent).
"""

import os
import sys
from pathlib import Path
from typing import Optional
from datetime import datetime, timedelta
import io

import streamlit as st
from docx import Document

# Add project root for imports
sys.path.insert(0, str(Path(__file__).resolve().parent))

from agents import AnalysisAgent, ReportAgent, SearchAgent


# -----------------------------------------------------------------------------
# Environment & config
# -----------------------------------------------------------------------------

def _load_env_file() -> None:
    """Optionally load ANTHROPIC_API_KEY from .env in project root (no extra deps)."""
    env_path = Path(__file__).resolve().parent / ".env"
    if not env_path.exists():
        return
    for line in env_path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        k, _, v = line.partition("=")
        k, v = k.strip(), v.strip().strip("'\"").strip()
        if k == "ANTHROPIC_API_KEY" and v:
            os.environ.setdefault("ANTHROPIC_API_KEY", v)
            break


def check_api_key() -> bool:
    """Verify ANTHROPIC_API_KEY is set. Show st.error and return False if missing."""
    key = os.environ.get("ANTHROPIC_API_KEY")
    if not key or not key.strip():
        st.error(
            "**ANTHROPIC_API_KEY is not set.**\n\n"
            "Set it in your environment before running the app, e.g.:\n\n"
            "```bash\n"
            "export ANTHROPIC_API_KEY='your-key-here'\n"
            "streamlit run app_demo.py\n"
            "```\n\n"
            "On Windows (PowerShell):\n\n"
            "```powershell\n"
            "$env:ANTHROPIC_API_KEY = 'your-key-here'\n"
            "streamlit run app_demo.py\n"
            "```"
        )
        return False
    return True


# -----------------------------------------------------------------------------
# Custom CSS
# -----------------------------------------------------------------------------

def inject_custom_css() -> None:
    """Apply professional, clinical SaaS styling via custom CSS."""
    st.markdown(
        """
        <style>
        /* Global layout */
        .main .block-container {
            padding-top: 1.2rem;
            padding-bottom: 2rem;
            padding-left: 2rem;
            padding-right: 2rem;
            max-width: 1200px;
        }

        /* Hide default Streamlit chrome */
        #MainMenu {visibility: hidden;}
        header[data-testid="stHeader"] {visibility: hidden;}
        footer {visibility: hidden;}

        /* App shell */
        .app-header {
            background: #020817;
            border-radius: 12px;
            border: 1px solid #1e293b;
            padding: 1.1rem 1.4rem;
            margin-bottom: 1.4rem;
            display: flex;
            align-items: center;
            justify-content: space-between;
            box-shadow: 0 12px 30px rgba(15, 23, 42, 0.4);
        }
        .app-header-left {
            display: flex;
            flex-direction: column;
            gap: 0.1rem;
        }
        .app-logo {
            width: 32px;
            height: 32px;
            border-radius: 10px;
            background: #0ea5e9;
            display: flex;
            align-items: center;
            justify-content: center;
            font-weight: 700;
            color: #0b1220;
            margin-right: 0.85rem;
            box-shadow: 0 6px 16px rgba(15, 23, 42, 0.45);
        }
        .app-title-row {
            display: flex;
            align-items: center;
            gap: 0.7rem;
            margin-bottom: 0.25rem;
        }
        .app-title {
            color: #e5e7eb;
            font-size: 1.2rem;
            font-weight: 600;
            letter-spacing: 0.02em;
        }
        .app-badge {
            font-size: 0.68rem;
            text-transform: uppercase;
            letter-spacing: 0.14em;
            border-radius: 999px;
            padding: 0.15rem 0.55rem;
            background: rgba(15,23,42,0.9);
            color: #9ca3af;
            border: 1px solid rgba(148,163,184,0.65);
        }
        .app-tagline {
            color: #9ca3af;
            font-size: 0.9rem;
            max-width: 460px;
        }
        .app-header-right {
            display: flex;
            flex-direction: column;
            align-items: flex-end;
            gap: 0.25rem;
        }
        .status-pill {
            font-size: 0.75rem;
            color: #e5e7eb;
            padding: 0.2rem 0.6rem;
            border-radius: 999px;
            border: 1px solid rgba(148,163,184,0.7);
            background: #020617;
        }
        .status-pill strong {
            color: #d1fae5;
        }
        .status-subtext {
            font-size: 0.78rem;
            color: #9ca3af;
        }

        /* Input panel */
        .input-card {
            background: #f9fafb;
            border-radius: 14px;
            padding: 1.3rem 1.35rem 1.1rem;
            border: 1px solid #e5e7eb;
            box-shadow: 0 10px 30px rgba(15, 23, 42, 0.04);
            margin-bottom: 1.4rem;
        }
        .input-card h3 {
            font-size: 1rem;
            font-weight: 600;
            color: #0f172a;
            margin-bottom: 0.4rem;
        }
        .input-description {
            font-size: 0.85rem;
            color: #6b7280;
            margin-bottom: 0.8rem;
        }

        /* Streamlit widgets */
        .stTextInput > div > div > input,
        .stNumberInput input {
            border-radius: 10px !important;
            border: 1px solid #d1d5db !important;
            padding: 0.4rem 0.6rem !important;
            font-size: 0.9rem !important;
        }
        .stTextInput > label,
        .stNumberInput > label {
            font-size: 0.85rem;
            font-weight: 500;
            color: #374151;
        }

        .stButton > button {
            border-radius: 999px;
            padding: 0.5rem 1.35rem;
            font-weight: 600;
            font-size: 0.9rem;
            border: 1px solid transparent;
            background: linear-gradient(135deg, #1d4ed8, #0f172a);
            color: #f9fafb;
            box-shadow: 0 10px 25px rgba(15, 23, 42, 0.22);
        }
        .stButton > button:hover {
            border-color: #93c5fd;
            box-shadow: 0 14px 30px rgba(15, 23, 42, 0.3);
        }

        /* Summary bar */
        .results-summary-bar {
            background: #0b1220;
            border-radius: 12px;
            padding: 0.7rem 0.85rem;
            display: flex;
            align-items: center;
            justify-content: space-between;
            margin-bottom: 0.9rem;
            border: 1px solid rgba(148, 163, 184, 0.3);
            box-shadow: 0 14px 30px rgba(15, 23, 42, 0.35);
            color: #e5e7eb;
        }
        .results-summary-left {
            display: flex;
            flex-direction: column;
            gap: 0.05rem;
        }
        .results-summary-title {
            font-size: 0.85rem;
            text-transform: uppercase;
            letter-spacing: 0.12em;
            color: #9ca3af;
        }
        .results-summary-count {
            font-size: 1rem;
            font-weight: 600;
        }
        .results-summary-right {
            display: flex;
            flex-wrap: wrap;
            gap: 0.4rem;
        }
        .summary-pill {
            font-size: 0.8rem;
            padding: 0.2rem 0.6rem;
            border-radius: 999px;
            background: rgba(15, 23, 42, 0.9);
            border: 1px solid rgba(148, 163, 184, 0.45);
        }
        .summary-pill span {
            font-weight: 600;
        }
        .summary-pill-high span { color: #bbf7d0; }
        .summary-pill-medium span { color: #fde68a; }
        .summary-pill-low span { color: #e5e7eb; }

        /* Result cards */
        .results-section-title {
            font-size: 0.9rem;
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 0.14em;
            color: #6b7280;
            margin-top: 1.3rem;
            margin-bottom: 0.4rem;
        }
        .result-card {
            background: #ffffff;
            border-radius: 10px;
            padding: 0.9rem 1rem 0.85rem;
            margin-bottom: 0.7rem;
            border: 1px solid #e5e7eb;
            box-shadow: 0 8px 18px rgba(15, 23, 42, 0.06);
        }
        .result-card-header {
            display: flex;
            justify-content: space-between;
            gap: 0.75rem;
            margin-bottom: 0.3rem;
        }
        .result-title {
            font-weight: 600;
            color: #111827;
            font-size: 0.96rem;
        }
        .relevance-badge {
            border-radius: 999px;
            padding: 0.15rem 0.55rem;
            font-size: 0.75rem;
            font-weight: 500;
        }
        .relevance-high {
            background: #ecfdf3;
            color: #166534;
            border: 1px solid #bbf7d0;
        }
        .relevance-medium {
            background: #fffbeb;
            color: #92400e;
            border: 1px solid #facc15;
        }
        .relevance-low {
            background: #f3f4f6;
            color: #374151;
            border: 1px solid #e5e7eb;
        }
        .result-meta {
            font-size: 0.8rem;
            color: #6b7280;
            margin-bottom: 0.25rem;
        }
        .result-section-label {
            font-size: 0.78rem;
            font-weight: 600;
            color: #4b5563;
            text-transform: uppercase;
            letter-spacing: 0.08em;
            margin-top: 0.35rem;
        }
        .result-text {
            font-size: 0.86rem;
            color: #374151;
        }
        .result-footer {
            margin-top: 0.45rem;
            display: flex;
            justify-content: space-between;
            align-items: center;
            gap: 0.5rem;
            font-size: 0.8rem;
        }
        .result-link a {
            color: #1d4ed8;
            text-decoration: none;
            font-weight: 500;
        }
        .result-link a:hover {
            text-decoration: underline;
        }
        .result-divider {
            margin-top: 0.7rem;
            border-top: 1px dashed #e5e7eb;
        }

        /* Download buttons */
        .download-bar {
            margin-top: 0.9rem;
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 0.6rem;
        }
        .download-help {
            font-size: 0.78rem;
            color: #6b7280;
        }
        [data-testid="stDownloadButton"] > button {
            border-radius: 999px;
            padding: 0.45rem 1.3rem;
            font-size: 0.88rem;
            font-weight: 600;
            border: 1px solid transparent;
            background: linear-gradient(135deg, #0ea5e9, #0369a1);
            color: #f9fafb;
            box-shadow: 0 10px 25px rgba(15, 23, 42, 0.18);
        }
        [data-testid="stDownloadButton"] > button:hover {
            border-color: #bae6fd;
            box-shadow: 0 14px 30px rgba(15, 23, 42, 0.28);
        }

        /* Status containers */
        [data-testid="stStatus"] {
            margin: 0.75rem 0;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


# -----------------------------------------------------------------------------
# Word report generation
# -----------------------------------------------------------------------------


def generate_word_report(publications: list[dict]) -> bytes:
    """Build a formatted .docx literature monitoring report in memory."""
    doc = Document()

    today_str = datetime.now().strftime("%Y-%m-%d")
    doc.add_heading(f"Literature Monitoring Report", level=0)
    doc.add_paragraph(f"Date generated: {today_str}")
    doc.add_paragraph()

    total = len(publications)
    high = sum(1 for p in publications if p.get("relevance") == "High")
    medium = sum(1 for p in publications if p.get("relevance") == "Medium")
    low = sum(1 for p in publications if p.get("relevance") == "Low")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(f"Total publications: {total}")
    doc.add_paragraph(f"High relevance: {high}")
    doc.add_paragraph(f"Medium relevance: {medium}")
    doc.add_paragraph(f"Low relevance: {low}")
    doc.add_paragraph()

    def add_section(title: str, label: str) -> None:
        section_pubs = [p for p in publications if p.get("relevance") == label]
        if not section_pubs:
            return

        doc.add_heading(title, level=1)
        for p in section_pubs:
            title_text = p.get("title") or "No title"
            summary_text = p.get("summary") or "No summary available."
            pmid = p.get("pmid") or ""
            url = p.get("url") or ""

            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title_text)
            title_run.bold = True

            doc.add_paragraph(f"Relevance: {label}")
            doc.add_paragraph(summary_text)

            if pmid and url:
                doc.add_paragraph(f"PMID: {pmid} ({url})")
            elif pmid:
                doc.add_paragraph(f"PMID: {pmid}")
            elif url:
                doc.add_paragraph(url)

            doc.add_paragraph("-" * 40)

        doc.add_paragraph()

    add_section("High Priority Publications", "High")
    add_section("Medium Priority Publications", "Medium")
    add_section("Low Priority Publications", "Low")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# -----------------------------------------------------------------------------
# Pipeline: search -> analyze -> report (unchanged core logic)
# -----------------------------------------------------------------------------

def _parse_publication_date(date_str: str) -> Optional[datetime]:
    """Best-effort parse of PubMed-style publication_date strings to a datetime."""
    if not date_str or date_str == "Unknown date":
        return None

    # Try a few common formats: YYYY, YYYY-MM, YYYY-MMM, YYYY-MM-DD
    for fmt in ("%Y-%m-%d", "%Y-%m", "%Y-%b", "%Y-%B", "%Y"):
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    return None


def run_pipeline(
    drug: str,
    days: int,
    max_results: int,
    therapeutic_area: Optional[str],
) -> tuple[list[dict], str, Optional[str]]:
    """
    Run the full pipeline: Search -> Analyze -> Report.

    Returns:
        (analyzed_publications, report_text, error_message).
        error_message is None on success.
    """
    try:
        # --- Phase 1: Search ---
        with st.status("**Searching PubMed...**", state="running", expanded=True) as search_status:
            search_agent = SearchAgent()
            publications = search_agent.search(
                drug_name=drug,
                therapeutic_area=therapeutic_area or None,
                days_back=days,
                max_results=max_results,
            )
            search_status.update(label="**Search complete.**", state="complete", expanded=False)

        # Enforce date range on returned publications as a safety net
        cutoff = datetime.now() - timedelta(days=days)
        filtered_publications: list[dict] = []
        for pub in publications:
            parsed = _parse_publication_date(str(pub.get("publication_date", "")))
            if parsed is None or parsed >= cutoff:
                filtered_publications.append(pub)

        publications = filtered_publications

        if not publications:
            return [], "", "No publications found. Try a longer date range or check the drug name."

        # --- Phase 2: Analysis (with progress) ---
        progress_bar = st.progress(0.0, text="Analyzing publications...")
        with st.status("**Analyzing publication 1 of {}...**".format(len(publications)), state="running", expanded=True) as analysis_status:

            def update_analysis_progress(current: int, total: int) -> None:
                progress_bar.progress(current / total, text=f"Analyzing publication {current} of {total}...")
                analysis_status.update(
                    label=f"**Analyzing publication {current} of {total}...**",
                    state="running",
                    expanded=True,
                )

            analysis_agent = AnalysisAgent(drug_name=drug, therapeutic_area=therapeutic_area or None)
            analyzed = analysis_agent.analyze_publications(publications, progress_callback=update_analysis_progress)

            analysis_status.update(
                label="**Analysis complete.**",
                state="complete",
                expanded=False,
            )
        progress_bar.progress(1.0, text="Analysis complete.")
        progress_bar.empty()

        # --- Phase 3: Report (formatted text for display + download) ---
        with st.status("**Generating report...**", state="running", expanded=True) as report_status:
            report_agent = ReportAgent(output_dir="reports")
            report_text = report_agent.get_report_content(
                publications=analyzed,
                drug_name=drug,
                therapeutic_area=therapeutic_area,
                days_back=days,
            )
            report_status.update(label="**Report generated.**", state="complete", expanded=False)

        return analyzed, report_text, None

    except ValueError as e:
        return [], "", str(e)
    except Exception as e:
        return [], "", f"An error occurred: {e}"


# -----------------------------------------------------------------------------
# UI: summary stats, expandable sections, download
# -----------------------------------------------------------------------------

def render_results(publications: list[dict], report_text: str, drug: str) -> None:
    """Render summary stats, publication cards, and download button."""

    def _count(relevance: str) -> int:
        return sum(1 for p in publications if p.get("relevance") == relevance)

    high, medium, low = _count("High"), _count("Medium"), _count("Low")

    # Results summary bar
    st.markdown(
        f"""
        <div class="results-summary-bar">
            <div class="results-summary-left">
                <div class="results-summary-title">Results overview</div>
                <div class="results-summary-count">{len(publications)} publications analyzed</div>
            </div>
            <div class="results-summary-right">
                <div class="summary-pill summary-pill-high"><span>{high}</span>&nbsp; High relevance</div>
                <div class="summary-pill summary-pill-medium"><span>{medium}</span>&nbsp; Medium relevance</div>
                <div class="summary-pill summary-pill-low"><span>{low}</span>&nbsp; Low relevance</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # Publication cards grouped by relevance
    def _render_priority_section(label: str, relevance: str) -> None:
        pubs = [p for p in publications if p.get("relevance") == relevance]
        if not pubs:
            return

        st.markdown(f'<div class="results-section-title">{label}</div>', unsafe_allow_html=True)
        for p in pubs:
            title = p.get("title") or "No title"
            summary = p.get("summary") or "No summary available."
            rationale = p.get("relevance_rationale") or "—"
            url = p.get("url") or ""
            journal = p.get("journal") or "Unknown"
            pub_date = p.get("publication_date") or "—"

            badge_class = {
                "High": "relevance-high",
                "Medium": "relevance-medium",
                "Low": "relevance-low",
            }.get(relevance, "relevance-low")

            st.markdown(
                f"""
                <div class="result-card">
                    <div class="result-card-header">
                        <div class="result-title">{title}</div>
                        <div class="relevance-badge {badge_class}">{relevance} relevance</div>
                    </div>
                    <div class="result-meta">{journal} · {pub_date}</div>
                    <div class="result-section-label">Summary</div>
                    <div class="result-text">{summary}</div>
                    <div class="result-section-label">Relevance rationale</div>
                    <div class="result-text">{rationale}</div>
                    <div class="result-footer">
                        <div class="result-link">
                            {"<a href='" + url + "' target='_blank'>View on PubMed →</a>" if url else ""}
                        </div>
                    </div>
                    <div class="result-divider"></div>
                </div>
                """,
                unsafe_allow_html=True,
            )

    _render_priority_section("High priority", "High")
    _render_priority_section("Medium priority", "Medium")
    _render_priority_section("Low priority", "Low")

    # Word document download
    safe_name = "".join(c if c.isalnum() else "_" for c in drug)
    report_bytes = generate_word_report(publications)
    today_str = datetime.now().strftime("%Y%m%d")
    filename = f"literature_report_{today_str}.docx"

    st.markdown(
        """
        <div class="download-bar">
            <div class="download-help">
                Export a structured Word report with priority tiers and individual publication details.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.download_button(
        label="Download Word report (.docx)",
        data=report_bytes,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# -----------------------------------------------------------------------------
# App entrypoint
# -----------------------------------------------------------------------------

def main() -> None:
    st.set_page_config(
        page_title="LIT Monitor",
        page_icon="📋",
        layout="wide",
        initial_sidebar_state="collapsed",
    )
    inject_custom_css()

    # Persist results across reruns (e.g. download button, expander toggle)
    if "report_data" not in st.session_state:
        st.session_state.report_data = None

    # Branded header (simplified, clinical)
    st.markdown(
        """
        <div class="app-header">
            <div class="app-header-left">
                <div class="app-title-row">
                    <div class="app-title">LIT Monitor</div>
                    <div class="app-badge">Medical affairs · Pharmacovigilance</div>
                </div>
                <div class="app-tagline">
                    Clinical literature monitoring for medical and safety teams.
                </div>
            </div>
            <div class="app-header-right">
                <div class="status-pill"><strong>Live</strong> · PubMed search & AI analysis</div>
                <div class="status-subtext">Searching PubMed → Analyzing publications → Generating report</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    _load_env_file()
    if not check_api_key():
        st.stop()

    # Input form: study configuration
    with st.container():
        st.markdown(
            """
            <div class="input-card">
                <h3>Configure your literature scan</h3>
                <div class="input-description">
                    Define the drug and clinical context. LIT Monitor will search PubMed, score relevance,
                    and return a structured report suitable for internal review.
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        with st.form("report_form"):
            drug = st.text_input(
                "Drug name",
                value="adalimumab",
                placeholder="e.g. adalimumab, pembrolizumab",
                help="Name of the drug to monitor.",
            ).strip()
            col_a, col_b, col_c = st.columns(3)
            with col_a:
                days = st.number_input(
                    "Date range (days)",
                    min_value=1,
                    max_value=365,
                    value=7,
                    help="Look back period for publications.",
                )
            with col_b:
                max_results = st.number_input(
                    "Max results",
                    min_value=5,
                    max_value=100,
                    value=50,
                    help="Maximum number of publications to fetch and analyze.",
                )
            with col_c:
                therapeutic_area = st.text_input(
                    "Therapeutic area (optional)",
                    value="",
                    placeholder="e.g. rheumatoid arthritis, oncology",
                ).strip() or None

            submitted = st.form_submit_button("Generate report")

    if submitted:
        if not drug:
            st.error("Please enter a drug name.")
        else:
            # Clear previous results before new run
            st.session_state.report_data = None
            analyzed, report_text, err = run_pipeline(
                drug, days, max_results, therapeutic_area
            )
            if err:
                st.error(err)
            else:
                st.session_state.report_data = (analyzed, report_text, drug)

    # Show cached results (after generate, or when user interacts with download/expanders)
    if st.session_state.report_data is not None:
        analyzed, report_text, drug = st.session_state.report_data
        render_results(analyzed, report_text, drug)
        st.divider()
        if st.button("Clear results and run new search"):
            st.session_state.report_data = None
            st.rerun()
    elif not submitted:
        st.info("Enter a drug name and click **Generate Report** to run the pipeline.")


if __name__ == "__main__":
    main()
